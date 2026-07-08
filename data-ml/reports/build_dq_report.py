#!/usr/bin/env python3
"""Data Quality Plan & Report (Venue OSM + MTA ridership) -> Word .docx"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- palette ----
NAVY   = RGBColor(0x0E, 0x1B, 0x2C)
GOLD   = RGBColor(0xB7, 0x82, 0x00)
GREEN  = RGBColor(0x2E, 0xA8, 0x6F)   # measured
AMBER  = RGBColor(0xC8, 0x7A, 0x12)   # inferred
RED    = RGBColor(0xC0, 0x39, 0x2B)   # synthetic
BLUE   = RGBColor(0x2B, 0x6C, 0xB0)   # computed / derived
GREY   = RGBColor(0x55, 0x5F, 0x6B)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

SHADE_HEAD = "0E1B2C"
SHADE_ALT  = "F2F4F7"
SHADE_GREEN= "E4F3EC"
SHADE_AMBER= "FBEFD9"
SHADE_RED  = "F8E3E0"
SHADE_BLUE = "E7EEF7"

doc = Document()

# base styles
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10)

for sec in doc.sections:
    sec.top_margin = Inches(0.7); sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.8); sec.right_margin = Inches(0.8)

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),hexcolor)
    tcPr.append(sh)

def set_cell(cell, text, bold=False, color=None, size=9, align='left', shade_hex=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {'left':WD_ALIGN_PARAGRAPH.LEFT,'center':WD_ALIGN_PARAGRAPH.CENTER,'right':WD_ALIGN_PARAGRAPH.RIGHT}[align]
    p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(1)
    r = p.add_run(str(text)); r.bold = bold; r.font.size = Pt(size)
    if color is not None: r.font.color.rgb = color
    if shade_hex: shade(cell, shade_hex)

def header_row(row):
    for c in row.cells:
        shade(c, SHADE_HEAD)
        for p in c.paragraphs:
            for r in p.runs:
                r.font.color.rgb = WHITE

def make_table(headers, rows, widths=None, prov_col=None, alt=True):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i,h in enumerate(headers):
        set_cell(hdr[i], h, bold=True, color=WHITE, size=9)
    header_row(t.rows[0])
    for ridx, row in enumerate(rows):
        cells = t.add_row().cells
        base_shade = SHADE_ALT if (alt and ridx % 2 == 1) else None
        for i,val in enumerate(row):
            sh = base_shade
            col = None
            # provenance coloring
            if prov_col is not None and i == prov_col:
                v = str(val).lower()
                if v.startswith("measured"): sh, col = SHADE_GREEN, GREEN
                elif v.startswith("inferred"): sh, col = SHADE_AMBER, AMBER
                elif v.startswith("synthetic"): sh, col = SHADE_RED, RED
                elif v.startswith("computed"): sh, col = SHADE_BLUE, BLUE
                set_cell(cells[i], val, bold=True, color=col, size=9, shade_hex=sh)
                continue
            set_cell(cells[i], val, size=9, shade_hex=sh)
    if widths:
        for i,w in enumerate(widths):
            for r in t.rows:
                r.cells[i].width = Inches(w)
    return t

def H1(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.bold=True; r.font.size=Pt(14); r.font.color.rgb = NAVY
    return p

def H2(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.bold=True; r.font.size=Pt(11.5); r.font.color.rgb = GOLD
    return p

def body(text, size=10, italic=False, color=None, space=4):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(space)
    r = p.add_run(text); r.font.size=Pt(size); r.italic=italic
    if color is not None: r.font.color.rgb = color
    return p

def bullet(text, size=10):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.font.size = Pt(size)
    return p

def caption(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(8)
    r = p.add_run(text); r.font.size=Pt(8.5); r.italic=True; r.font.color.rgb=GREY

def part(label, subtitle):
    """Full-width dark banner (single-cell table) marking a major part of the report."""
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    cell.width = Inches(10.5)
    shade(cell, SHADE_HEAD)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label.upper()); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=WHITE
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run(subtitle); r2.font.size=Pt(9.5); r2.italic=True
    r2.font.color.rgb = RGBColor(0xC7, 0xD2, 0xDC)
    return t

# ============================ TITLE ============================
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0)
r = p.add_run("Plug & WiFi  ·  Data / ML"); r.font.size=Pt(9); r.bold=True; r.font.color.rgb=GOLD
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0)
r = p.add_run("Data Quality Plan & Report"); r.font.size=Pt(22); r.bold=True; r.font.color.rgb=NAVY
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)
r = p.add_run("Profiling, remediation and enrichment of the two source datasets behind the busyness model: the "
              "OpenStreetMap venue table and the MTA subway ridership extract.")
r.font.size=Pt(10.5); r.italic=True; r.font.color.rgb=GREY
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(8)
r = p.add_run("fetch_venues.py + fetch_turnstiles.py  →  enrichment / aggregation    ·    Adam T    ·    1 July 2026"); r.font.size=Pt(9); r.font.color.rgb=GREY

# provenance legend
H2("Provenance legend")
leg = make_table(
    ["Tag","Meaning"],
    [
        ["Measured","Observed from a real-world source (OSM tags or computed geometry). Trustworthy."],
        ["Inferred","Deterministic rule from venue type. Reasonable default, not observed per-venue."],
        ["Synthetic","Random placeholder standing in for data we don't yet have. Do NOT treat as real."],
        ["Computed","Deterministic derivation of measured values (e.g. an average). As trustworthy as its inputs."],
    ],
    widths=[1.2, 9.3], prov_col=0, alt=False)
caption("Every field below is tagged with one of these. The distinction is the core message of this report.")

# ============================ PART A ============================
part("Part A · Venue data — OpenStreetMap",
     "One row per venue. 13,863 venues, 17 raw fields enriched to 35 columns. Feeds the work-friendliness score.")

# ============================ 1. OVERVIEW ============================
H1("1 · Dataset overview")
body("Source: OpenStreetMap Overpass API (NYC bounding box). One row per venue. The raw fetch pulls 17 fields; "
     "after enrichment and normalisation the table carries 35 columns. All figures below are computed on the current "
     "extract — 13,863 venues.")
ov = make_table(
    ["Metric","Value"],
    [
        ["Venues (rows)","13,863"],
        ["Columns — raw fetch / final","17 / 35"],
        ["Venue types","restaurant 9,464 · cafe 2,842 · hotel 796 · bakery 761"],
        ["Boroughs","Manhattan 6,925 · Queens 2,980 · Bronx 1,468 · Brooklyn 1,458 · Staten Island 1,032"],
        ["Geographic coverage","lat/lon present for 100% of rows; borough assigned by lat/lon rule"],
    ],
    widths=[2.3, 8.2], alt=True)

# ============================ 2. ORIGINAL FEATURES ============================
H1("2 · Original features pulled from fetch_venues.py")
body("These are the fields as they arrive from OSM, before any enrichment. The headline issue is sparse optional "
     "tags: contact details, address parts and opening hours are missing for roughly a third to over half of venues, "
     "and a usable Wi-Fi tag exists for only 6.3% of rows.")
rows = [
    ["venue_id","ID","Measured","0.0%","OSM element id, prefixed osm_; unique key"],
    ["name","Text","Measured","0.0%","Falls back to literal \"Unnamed\" when the OSM name tag is absent"],
    ["osm_type / cuisine_type","Categorical","Measured","0.0%","cafe / hotel / bakery / restaurant (derived from the OSM tag queried)"],
    ["cuisine_detail","Text","Measured","34.6%","Free-text OSM cuisine tag (e.g. italian, coffee_shop)"],
    ["phone","Text","Measured","49.8%","phone / contact:phone"],
    ["website","Text","Measured","53.6%","website / contact:website"],
    ["building_number","Text","Measured","30.4%","addr:housenumber"],
    ["street","Text","Measured","29.5%","addr:street"],
    ["zipcode","Text","Measured","45.1%","addr:postcode"],
    ["lat / lon","Float","Measured","0.0%","Node coords, or way centroid via out center"],
    ["opening_hours","Text","Measured","58.4%","Raw OSM opening_hours string (unparsed)"],
    ["has_wifi","Boolean","Measured","93.7%","From internet_access tag — only 795 yes / 85 no; the rest unknown"],
    ["best_hours_for_work","JSON","—","100.0%","Empty at fetch; populated later by noise_model"],
    ["hourly_profile","JSON","—","100.0%","Empty at fetch; populated later by noise_model"],
    ["partner","Boolean","Measured","0.0%","App-owned flag; all False in this extract"],
]
make_table(["Field","Type","Provenance","Missing %","Notes"], rows,
           widths=[2.1, 1.0, 1.2, 0.9, 5.3], prov_col=2, alt=True)
caption("Missing % treats empty strings and None as missing. has_wifi missingness is the single largest data gap in the raw pull.")

# ============================ 3. QUALITY ISSUES ============================
H1("3 · Data-quality issues identified")
bullet("Sparse Wi-Fi signal — only 880 of 13,863 venues (6.3%) carry any internet_access tag; 93.7% unknown.")
bullet("Incomplete addresses — street/number/zip missing for ~30–45% of venues; not all rows are geocodable to a postal address (lat/lon are always present, so spatial work is unaffected).")
bullet("Missing opening hours — 58.4% absent, and present values are raw OSM strings needing parsing before use.")
bullet("Name fallback — venues with no OSM name become \"Unnamed\"; needs handling before display.")
bullet("Borough by heuristic — assigned from lat/lon cut-offs, not an authoritative boundary file; edge cases near borough lines may misclassify.")
bullet("No ground-truth labels for the work-friendliness signals (Wi-Fi quality, plugs, noise, rating, price) — these do not exist in OSM and had to be supplied by the enrichment layer.")

# ============================ 4. ENRICHMENT & REMEDIATION ============================
H1("4 · Remediation & enrichment")
body("Each enrichment model writes one or more work-friendliness signals onto the venue table. The table below states, "
     "honestly, how each was produced. Three of the six are synthetic placeholders — usable for wiring up the app end-to-end, "
     "but they must be replaced with real sources before the signals are trusted.")
rows = [
    ["wifi_model","inferred_wifi, wifi_user_reported","Use the real OSM tag where present (6.3%); otherwise infer by type — cafe & hotel = yes, restaurant & bakery = unknown. User-reported column reserved to override later.","Measured + Inferred","3,706 yes · 85 no · 10,072 still unknown"],
    ["transit_model","nearest_subway(_m), nearest_bus(_m)","Haversine distance via a KD-tree to OSM subway stations and bus stops. Pure geometry, no guessing.","Measured","Subway: mean 1,454 m / median 312 m. Bus: mean 141 m / median 71 m"],
    ["noise_model","hourly_profile, best_hours_for_work","Rule-based score = base level (by type) × hour-of-day factor × weekend factor, bucketed quiet/moderate/loud. Deterministic from venue type.","Inferred","100% coverage; 24-hour profile per venue"],
    ["plug_model","plug_access, plug_user_reported","Random Bernoulli draw with a per-type probability (hotel 0.70, cafe 0.35, restaurant 0.10, bakery 0.05). Placeholder.","Synthetic","18.0% flagged as having plugs"],
    ["rating_model","rating, rating_user_reported","Random draw N(3.5, 1) clipped to 0–5. Placeholder for a real reviews feed.","Synthetic","mean 3.48 · std 0.94"],
    ["pricing_model","hourly_price, actual_hourly_price","Random draw N(9, 4) clipped at 0. Placeholder for real partner pricing.","Synthetic","mean $9.01 · std $3.92"],
]
make_table(["Model","Output field(s)","Method","Provenance","Result / key stat"], rows,
           widths=[1.4, 1.9, 4.2, 1.2, 1.8], prov_col=3, alt=True)
caption("'User-reported' companion columns (wifi_user_reported, plug_user_reported, rating_user_reported, actual_hourly_price) are intentionally empty now — they are the slots where real, app-collected data will later override the inferred/synthetic values.")

# ============================ 5. NORMALISATION ============================
H1("5 · Normalisation — model-ready features")
body("normalise.py converts the enriched signals into 0–1 features the scoring system can combine on a common scale. "
     "All five are fully populated (0% missing) because the normaliser fills gaps explicitly (e.g. unknown Wi-Fi → 0.5).")
rows = [
    ["plug_norm","plug_access as float (0/1)","0 – 1","0.180","Inherits the synthetic plug placeholder"],
    ["wifi_norm","inferred_wifi as float, unknown → 0.5","0 – 1","0.631","Neutral 0.5 fill is why mean sits above the 'yes' rate"],
    ["rating_norm","rating ÷ 5","0 – 1","0.697","Inherits the synthetic rating placeholder"],
    ["bus_norm","(max − dist) / (max − min), dist capped at 1,200 m","0 – 1","0.895","Closer bus stop → higher score; most venues very close"],
    ["train_norm","(max − dist) / (max − min), dist capped at 1,200 m","0 – 1","0.600","Closer subway → higher score"],
]
make_table(["Feature","Formula","Range","Mean","Note"], rows,
           widths=[1.3, 4.0, 1.0, 0.9, 3.3], alt=True)
caption("Distance features are inverted and min-max scaled so that 'nearer = better'. The 1,200 m cap stops far-flung "
        "outliers (subway max 20.5 km) from compressing the useful range.")

# ============================ 6. PROVENANCE SUMMARY ============================
H1("6 · Provenance summary")
body("Of the work-friendliness signals layered onto each venue:")
bullet("Measured (trust): coordinates, venue type, transit distances, and the 6.3% of Wi-Fi tags that came from OSM.")
bullet("Inferred (reasonable default): the noise profile and the type-based Wi-Fi fill.")
bullet("Synthetic (replace before trusting): rating, hourly_price, and plug_access are random placeholders. They let the "
       "end-to-end app and scoring pipeline run today, but carry no real-world signal.")

# ============================ PART B ============================
part("Part B · MTA subway ridership",
     "Hourly ridership per station, the real-world demand signal the busyness model is trained to predict.")

# ---- B1 overview ----
H1("7 · Dataset overview")
body("Source: MTA Subway Hourly Ridership via the NYC Open Data Socrata API (dataset 5wq4-mkjj). "
     "Unlike the venue pull, the heavy lifting happens server-side: fetch_turnstiles.py asks the API to GROUP BY "
     "station, hour and day-type and to SUM ridership, so what comes back is already aggregated and clean. "
     "The grain is one row per station × hour-of-day (0–23) × weekday/weekend — about 48 rows per station.")
ov2 = make_table(
    ["Metric","Value"],
    [
        ["Rows","20,542"],
        ["Columns","10"],
        ["Stations","428"],
        ["Grain","station × hour (0–23) × weekday/weekend  ≈ 48 rows/station"],
        ["Boroughs","Brooklyn 7,488 · Manhattan 5,902 · Queens 3,792 · Bronx 3,264 · Staten Island 96"],
        ["Day-type split","weekday 10,271 · weekend 10,271 (balanced)"],
        ["Window","from 2026-03-11 (≈71 weekday days · ≈28 weekend days)"],
        ["Missing values","0% on every column — a consequence of server-side aggregation"],
    ],
    widths=[2.3, 8.2], alt=True)

# ---- B2 fields ----
H1("8 · Fields in the ridership extract")
body("All ten columns arrive from the API. Five describe the station (measured), one is the hour bucket, and the "
     "remaining columns are either summed or derived by the query itself — hence tagged Computed rather than Measured.")
rows = [
    ["station_complex_id","ID","Measured","0.0%","Station complex key; groups platforms at one interchange"],
    ["station_complex","Text","Measured","0.0%","Human-readable station name"],
    ["borough","Categorical","Measured","0.0%","Bk / M / Q / Bx / SI"],
    ["latitude / longitude","Float","Measured","0.0%","Station coordinates — used to match venues to nearby stations"],
    ["hour","Integer","Measured","0.0%","Hour of day 0–23 (date_extract_hh on the timestamp)"],
    ["total_ridership","Integer","Computed","0.0%","SUM of ridership over the window for that station/hour/day-type. Mean 18,370 · median 5,449 · max 1,308,067 (right-skewed)"],
    ["day_type","Categorical","Computed","0.0%","weekday / weekend, from a CASE on day-of-week"],
    ["num_days","Integer","Computed","0.0%","Distinct dates in the window for that day-type. Weekday ≈71 · weekend ≈28"],
    ["avg_ridership","Float","Computed","0.0%","total_ridership ÷ num_days — the per-day figure the model targets. Mean 337 · median 131 · max 18,423"],
]
make_table(["Field","Type","Provenance","Missing %","Notes"], rows,
           widths=[2.1, 1.0, 1.2, 0.9, 5.3], prov_col=2, alt=True)
caption("avg_ridership is the modelling target. It is Computed, not Measured — but its inputs (summed swipes, counted days) are real, so it is trustworthy.")

# ---- B3 quality notes ----
H1("9 · Data-quality notes — ridership")
bullet("Clean by construction — 0% missing on all columns because the API returns pre-aggregated groups; there are no raw per-swipe rows to carry gaps.")
bullet("Right-skew → log transform — avg_ridership has mean 337 but median 131 and a max of 18,423 (std ≈692). A handful of hub stations dominate, so the model works on log ridership rather than the raw value.")
bullet("Average, not total — weekday and weekend windows differ (~71 vs ~28 days), so raw totals aren't comparable across day-types. Dividing by num_days puts every row on a per-day footing.")
bullet("Consistent grain — ~48 rows per station (24 hours × 2 day-types). A few stations have 46–47 where a particular hour saw zero ridership and dropped out of the group.")
bullet("Staten Island is thin — only 2 SI stations (96 rows). The Staten Island Railway is largely outside this subway feed, so venue-to-station matching there rests on very few points.")

# ============================ 10. RECOMMENDATIONS ============================
H1("10 · Recommended next steps")
rows = [
    ["Replace synthetic rating","High","Venue","Wire in a real reviews source (a TripAdvisor scraper already exists in the repo) → fill rating_user_reported / rating."],
    ["Replace synthetic price","High","Venue","Source hourly_price from partner agreements; keep actual_hourly_price for confirmed values."],
    ["Replace synthetic plug data","High","Venue","Seed from partner data; collect plug_user_reported in-app to override."],
    ["Backfill Wi-Fi","Medium","Venue","10,072 venues (mostly restaurants/bakeries) have unknown Wi-Fi — collect via user reports or a venue feed."],
    ["Validate noise heuristic","Medium","Venue","Sanity-check the rule-based noise profile against any real observation before it drives the indicator."],
    ["Widen the ridership window","Medium","Ridership","The current window starts 2026-03-11; a longer span steadies per-hour averages and softens one-off events."],
    ["Improve Staten Island coverage","Low","Ridership","Only 2 SI subway stations — consider a Staten Island Railway or bus feed so SI venues aren't matched on sparse data."],
    ["Harden borough + names","Low","Venue","Replace lat/lon borough rule with a boundary file; resolve \"Unnamed\" venues for display."],
    ["Parse opening_hours","Low","Venue","Convert raw OSM strings into structured open/closed windows."],
]
make_table(["Action","Priority","Dataset","Detail"], rows, widths=[2.2, 0.8, 0.9, 6.6], alt=True)

# ============================ 11. MODEL PERFORMANCE ============================
H1("11 · Model performance — busyness KNN")
body("The busyness model predicts subway ridership from location and time, then reads it out at venue "
     "coordinates. Errors are measured on the honest split (20% of stations held out in full) and reported "
     "in riders per day — predictions are exponentiated back from log space before scoring. Chosen model: "
     "distance-weighted KNN, k = 10.")
rows = [
    ["MAE — mean absolute error","217.8","Average size of a prediction miss, in riders/day. The headline, most interpretable figure."],
    ["RMSE — root mean squared error","682.0","Squares errors before averaging, so large misses dominate. Here RMSE is ~3× MAE, signalling a heavy-tailed error distribution."],
]
make_table(["Metric","Value (riders/day)","What it measures"], rows, widths=[2.7, 1.7, 6.1], alt=True)
caption("The wide RMSE-to-MAE gap is the Part B skew resurfacing: absolute error grows with ridership, so a few very busy hub stations are hard to hit in raw counts even when the log-scale fit is sound.")

body("Two metrics don't fully characterise a skewed target. Worth adding before this is presented as final:")
bullet("Median absolute error (MedAE) — the typical miss for an ordinary station, unmoved by the handful of giants. It will sit well below the 217.8 mean and better reflects the experience at most venues.")
bullet("R² on the grouped split — variance explained versus simply predicting the mean. It puts the raw error in context and lets KNN and the Random Forest be compared on one number (this is the honest replacement for the old, leaky OOB R² of 0.93).")
bullet("Baseline-relative MAE — quote the error against the k = 1 'nearest single station' baseline, so the gain from averaging over neighbours is explicit rather than assumed.")
bullet("Log-scale error (e.g. RMSLE) — the model is fitted on log ridership, so a log-scale error counts a 2× miss equally at a small station and a hub, matching how the model actually treats the problem.")

doc.save("/Users/adamt/Documents/workshare/data_quality_report.docx")
print("saved docx with", len(doc.tables), "tables")
